"""Round-trip tests for the docx template editor, on a synthetic document.

The user's real templates hold personal data and stay out of git, so these
tests build a minimal document with the same structure: an ALL-CAPS experience
section, ``Title | Company\\tDates`` headings, List-styled bullets, and a
table-based letter with subject, body paragraphs, and a German closing.
"""

import pytest
from docx import Document

from app.services.docx_template import (
    apply_anschreiben,
    apply_cv_bullets,
    parse_anschreiben,
    parse_cv_experiences,
    set_paragraph_text,
)


def _make_cv() -> Document:
    doc = Document()
    doc.add_paragraph("JANE EXAMPLE")
    doc.add_paragraph("BERUFSERFAHRUNG")
    doc.add_paragraph("Engineer | AcmeCo\tAug 2024 – Ongoing")
    doc.add_paragraph("Berlin, Germany")
    for i in range(2):
        doc.add_paragraph(f"Old Acme bullet {i}", style="List Bullet")
    doc.add_paragraph("Developer | DataCo\tSep 2023 – Jul 2024")
    doc.add_paragraph("Old Data bullet 0", style="List Bullet")
    doc.add_paragraph("STUDIUM")
    doc.add_paragraph("B.Sc. Somewhere\t2019 – 2023")
    return doc


def _make_letter() -> Document:
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.paragraphs[0].text = "17 Juni 2026"
    cell.add_paragraph("Alter Betreff der Bewerbung als Beispiel Rolle")
    cell.add_paragraph("Erster alter Absatz.")
    cell.add_paragraph("Zweiter alter Absatz.")
    # Decomposed umlaut (u + combining diaeresis), as Word often stores it.
    cell.add_paragraph("Mit freundlichen Grüßen,")
    cell.add_paragraph("Jane Example")
    return doc


def test_parse_cv_finds_experiences_and_bullets():
    exps = parse_cv_experiences(_make_cv())
    assert [(e.title, e.company, len(e.bullet_paras)) for e in exps] == [
        ("Engineer", "AcmeCo", 2),
        ("Developer", "DataCo", 1),
    ]
    assert exps[0].dates == "Aug 2024 – Ongoing"


def test_apply_cv_bullets_replaces_only_bullet_text():
    doc = _make_cv()
    exps = parse_cv_experiences(doc)
    apply_cv_bullets(exps, [["New A0", "New A1"], ["New D0"]])

    reparsed = parse_cv_experiences(doc)
    assert reparsed[0].bullets == ["New A0", "New A1"]
    assert reparsed[1].bullets == ["New D0"]
    # Headings, location line, and sections are untouched.
    texts = [p.text for p in doc.paragraphs]
    assert "Engineer | AcmeCo\tAug 2024 – Ongoing" in texts
    assert "Berlin, Germany" in texts
    assert "STUDIUM" in texts


def test_apply_cv_bullets_rejects_count_mismatch():
    doc = _make_cv()
    exps = parse_cv_experiences(doc)
    with pytest.raises(ValueError, match="counts must match"):
        apply_cv_bullets(exps, [["only one"], ["New D0"]])


def test_parse_and_apply_anschreiben_with_decomposed_umlauts():
    doc = _make_letter()
    letter = parse_anschreiben(doc)
    assert letter is not None
    assert letter.subject_para.text.startswith("Alter Betreff")
    assert len(letter.body_paras) == 2

    apply_anschreiben(letter, ["Neuer erster Absatz.", "Neuer zweiter Absatz."],
                      "Bewerbung als Backend Engineer (m/w/d)")
    texts = [p.text for p in letter.cell.paragraphs]
    assert "Neuer erster Absatz." in texts
    assert "Bewerbung als Backend Engineer (m/w/d)" in texts
    # Closing and signature untouched.
    assert any("Grüßen" in t or "Grüßen" in t for t in texts)
    assert "Jane Example" in texts


def test_set_paragraph_text_preserves_run_count_shape():
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("old")
    run.bold = True
    p.add_run(" tail")
    set_paragraph_text(p, "new text")
    assert p.text == "new text"
    assert p.runs[0].bold is True  # formatting of the first run survives


# --- filename collision guard --------------------------------------------


def test_role_slug_disambiguates_two_roles_at_one_company():
    from app.services.docx_prepare import _role_slug

    a = _role_slug("Senior Software Developer (m/w/d) - Java")
    b = _role_slug("Werkstudent Machine Learning")
    assert a and b and a != b
    # Filesystem-safe: no punctuation or spaces survive.
    for slug in (a, b):
        assert all(c.isalnum() or c == "_" for c in slug)


def test_role_slug_empty_without_a_title():
    from app.services.docx_prepare import _role_slug

    assert _role_slug(None) == ""
    assert _role_slug("") == ""
