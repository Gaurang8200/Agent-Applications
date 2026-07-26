"""Prepare an application from the user's own .docx templates.

Source of truth is the user's real CV and Anschreiben files. The flow:

1. Parse the CV docx: experiences with their existing bullets, and the skills
   section. Parse the Anschreiben docx: subject slot + body paragraph slots.
2. Tailor with Claude, grounded in exactly that parsed content, with bullet
   counts and paragraph counts taken from the documents themselves.
3. Write the tailored text back INTO copies of the user's files — only bullet
   text, the letter subject, and the letter body change. Name, contact, dates,
   layout, and design are untouched by construction.
4. Convert both to PDF with LibreOffice headless and save everything under
   APPLICATIONS_DIR/<Company>/.
"""

import shutil
import subprocess
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document

from app.core.config import get_settings
from app.models import Profile, Skill, WorkExperience
from app.schemas.tailor import TailoringConstraints, TailorResult
from app.services.application_files import safe_company_name
from app.services.docx_template import (
    apply_anschreiben,
    apply_cv_bullets,
    collect_cv_skills,
    parse_anschreiben,
    parse_cv_experiences,
)

settings = get_settings()


class DocxPrepareError(RuntimeError):
    pass


@dataclass
class PreparedFiles:
    company: str
    cv_docx: str
    anschreiben_docx: str
    cv_pdf: str
    anschreiben_pdf: str


def _find_soffice() -> str:
    candidates = [
        settings.soffice_path,
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        shutil.which("soffice") or "",
        shutil.which("libreoffice") or "",
    ]
    for c in candidates:
        if c and Path(c).exists():
            return c
    raise DocxPrepareError(
        "LibreOffice not found — needed to convert .docx to PDF. Install it "
        "(brew install --cask libreoffice) or set SOFFICE_PATH."
    )


def _parse_date(raw: str) -> date | None:
    # Dates in the CV are display strings ("Aug 2024"); the tailor prompt only
    # echoes them, so parsing precision doesn't matter. Keep None.
    return None


def build_profile_from_docx(cv_doc: Document) -> tuple[Profile, list[int]]:
    """An in-memory Profile built purely from the user's CV document.

    Returned alongside the per-experience bullet counts, which drive both the
    tailor constraints and the write-back (counts must match exactly).
    """
    experiences = parse_cv_experiences(cv_doc)
    if not experiences:
        raise DocxPrepareError(
            "No experiences found under BERUFSERFAHRUNG in the CV template."
        )

    profile = Profile()
    for order, exp in enumerate(experiences):
        profile.work_experience.append(
            WorkExperience(
                company=exp.company,
                title=exp.title,
                display_order=order,
                is_current="ongoing" in exp.dates.lower() or "heute" in exp.dates.lower(),
                highlights=exp.bullets,
            )
        )
    for name in collect_cv_skills(cv_doc):
        profile.skills.append(Skill(name=name))

    counts = [len(exp.bullet_paras) for exp in experiences]
    return profile, counts


def constraints_from_templates(
    cv_doc: Document, letter_paragraphs: int
) -> TailoringConstraints:
    """Derive formatting constraints from the user's own document.

    The chars-per-line estimate comes from the template itself: the longest
    existing bullet is taken to fill two rendered lines. The band then targets
    1.8–1.97 lines, inside the spec's 1.75–2.0 with a safety margin on both
    edges (character counts are a proxy — German compound words shift wrap
    points).
    """
    experiences = parse_cv_experiences(cv_doc)
    counts = [len(exp.bullet_paras) for exp in experiences]

    all_bullets = [b for exp in experiences for b in exp.bullets if b]
    longest = max((len(b) for b in all_bullets), default=210)
    chars_per_line = max(longest / 2, 80.0)

    return TailoringConstraints(
        bullets_per_experience=counts,
        bullet_min_chars=int(chars_per_line * 1.80),
        bullet_max_chars=int(chars_per_line * 1.97),
        anschreiben_paragraphs=letter_paragraphs,
    )


def _convert_pdf(soffice: str, docx_path: Path, outdir: Path) -> Path:
    result = subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(outdir), str(docx_path)],
        capture_output=True,
        text=True,
        timeout=180,
    )
    pdf_path = outdir / (docx_path.stem + ".pdf")
    if result.returncode != 0 or not pdf_path.exists():
        raise DocxPrepareError(
            f"PDF conversion failed for {docx_path.name}: "
            f"{result.stderr.strip() or result.stdout.strip()}"
        )
    return pdf_path


def write_tailored_files(
    *,
    cv_template: Path,
    anschreiben_template: Path,
    result: TailorResult,
    company: str,
) -> PreparedFiles:
    """Apply the tailored text into copies of the user's templates and render PDFs."""
    safe = safe_company_name(company)
    outdir = settings.applications_path / safe
    outdir.mkdir(parents=True, exist_ok=True)

    cv_out = outdir / f"CV_{safe}.docx".replace(" ", "_")
    letter_out = outdir / f"Anschreiben_{safe}.docx".replace(" ", "_")
    shutil.copyfile(cv_template, cv_out)
    shutil.copyfile(anschreiben_template, letter_out)

    # CV: swap only bullet text.
    cv_doc = Document(str(cv_out))
    experiences = parse_cv_experiences(cv_doc)
    new_bullets = [exp.bullets for exp in result.cv.experiences]
    apply_cv_bullets(experiences, new_bullets)
    cv_doc.save(str(cv_out))

    # Anschreiben: swap subject + body paragraphs.
    letter_doc = Document(str(letter_out))
    letter = parse_anschreiben(letter_doc)
    if letter is None:
        raise DocxPrepareError("Could not locate the letter body in the Anschreiben template.")
    apply_anschreiben(letter, result.anschreiben.paragraphs, result.anschreiben.subject)
    letter_doc.save(str(letter_out))

    soffice = _find_soffice()
    cv_pdf = _convert_pdf(soffice, cv_out, outdir)
    letter_pdf = _convert_pdf(soffice, letter_out, outdir)

    return PreparedFiles(
        company=company,
        cv_docx=str(cv_out),
        anschreiben_docx=str(letter_out),
        cv_pdf=str(cv_pdf),
        anschreiben_pdf=str(letter_pdf),
    )


def prepare_application(
    *,
    cv_template: Path,
    anschreiben_template: Path,
    job_description: str,
    job_title: str | None,
    company: str,
    max_render_retries: int = 1,
) -> tuple[PreparedFiles, "TailorResult", list[str]]:
    """Full docx flow: tailor -> write into the templates -> PDF -> verify lines.

    The rendered PDF is measured against the 1.75–2.0-line bullet rule. If any
    bullet misses, the measured problems are fed back to the tailor stage once
    and the documents are re-rendered. Returns the files, the final tailor
    result, and any line problems that remain (empty = fully verified).
    """
    from app.agents.tailor import tailor  # local import to avoid cycle
    from app.services.pdf_lines import line_violations, measure_bullets

    cv_doc = Document(str(cv_template))
    letter = parse_anschreiben(Document(str(anschreiben_template)))
    if letter is None:
        raise DocxPrepareError(
            "Could not locate the letter body in the Anschreiben template."
        )

    profile, _ = build_profile_from_docx(cv_doc)
    constraints = constraints_from_templates(cv_doc, len(letter.body_paras))
    sample_body = "\n\n".join(p.text for p in letter.body_paras)

    correction: str | None = None
    files = result = None
    problems: list[str] = []
    for _attempt in range(max_render_retries + 1):
        result = tailor(
            profile,
            job_description,
            job_title=job_title,
            company=company,
            constraints=constraints,
            sample_anschreiben=sample_body,
            extra_correction=correction,
        )
        files = write_tailored_files(
            cv_template=cv_template,
            anschreiben_template=anschreiben_template,
            result=result,
            company=company,
        )
        bullets = [b for exp in result.cv.experiences for b in exp.bullets]
        problems = line_violations(measure_bullets(files.cv_pdf, bullets))
        if not problems:
            break
        correction = (
            "The rendered PDF shows these line-length problems. Every bullet "
            "must fill between 1.75 and 2.0 rendered lines:\n"
            + "\n".join(f"- {p}" for p in problems)
        )

    return files, result, problems
