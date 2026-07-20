import io

from pypdf import PdfReader

SUPPORTED_CONTENT_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
}

MAX_UPLOAD_BYTES = 10 * 1024 * 1024


class UnsupportedFileType(ValueError):
    pass


class TextExtractionError(RuntimeError):
    pass


def extract_text(data: bytes, content_type: str) -> str:
    kind = SUPPORTED_CONTENT_TYPES.get(content_type)
    if kind is None:
        raise UnsupportedFileType(f"Unsupported content type: {content_type}")

    if kind == "pdf":
        return _extract_pdf(data)
    if kind == "docx":
        return _extract_docx(data)
    return data.decode("utf-8", errors="replace")


def _extract_pdf(data: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # pypdf raises a variety of parse errors
        raise TextExtractionError(f"Could not read PDF: {exc}") from exc

    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n\n".join(pages).strip()
    if not text:
        # Almost always a scanned/image-only PDF. OCR would be the fix; we do
        # not silently return an empty profile.
        raise TextExtractionError(
            "No selectable text found. If this is a scanned resume, upload a "
            "text-based PDF or DOCX instead."
        )
    return text


def _extract_docx(data: bytes) -> str:
    from docx import Document  # imported lazily; python-docx is slow to import

    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:
        raise TextExtractionError(f"Could not read DOCX: {exc}") from exc

    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts).strip()
    if not text:
        raise TextExtractionError("DOCX contained no readable text.")
    return text
