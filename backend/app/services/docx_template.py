"""Edit the candidate's own CV and Anschreiben .docx in place.

The whole point: keep the user's exact design, fonts, header, contact details,
and layout untouched, and swap only the experience bullet text (CV) and the
letter body text (Anschreiben). Nothing else in the document changes.

We never rebuild the document from a template of our own — we open theirs and
rewrite specific paragraphs, preserving each paragraph's run formatting.
"""

import re
import unicodedata
from dataclasses import dataclass, field

from docx import Document
from docx.table import _Cell
from docx.text.paragraph import Paragraph

# Section headings in the CV are short ALL-CAPS lines (BERUFSERFAHRUNG, STUDIUM,
# KENNTNISSE, SPRACHE). Used to bound the experience block.
_SECTION_HEADING = re.compile(r"^[A-ZÄÖÜ][A-ZÄÖÜ ]{3,}$")


@dataclass
class ParsedExperience:
    title: str
    company: str
    dates: str
    heading_text: str
    bullet_paras: list[Paragraph] = field(default_factory=list)

    @property
    def bullets(self) -> list[str]:
        return [p.text.strip() for p in self.bullet_paras]


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """Replace a paragraph's text while keeping its formatting.

    The first run keeps its font/size/bold and takes the new text; any further
    runs are emptied (not deleted, so numbering/bullet properties survive).
    """
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _is_bullet(paragraph: Paragraph) -> bool:
    return paragraph.style is not None and "List" in (paragraph.style.name or "")


def parse_cv_experiences(doc: Document) -> list[ParsedExperience]:
    """Find each job in the experience section with its bullet paragraphs.

    A job heading is a non-bullet paragraph containing a pipe (``Title | Company``)
    inside the experience section. Bullets are the ``List`` styled paragraphs that
    follow it, up to the next heading or the end of the section.
    """
    paras = doc.paragraphs
    # Locate the experience section bounds.
    start = None
    for i, p in enumerate(paras):
        if p.text.strip().upper() == "BERUFSERFAHRUNG":
            start = i + 1
            break
    if start is None:
        return []
    end = len(paras)
    for i in range(start, len(paras)):
        t = paras[i].text.strip()
        if t and _SECTION_HEADING.match(t) and "|" not in t:
            end = i
            break

    experiences: list[ParsedExperience] = []
    current: ParsedExperience | None = None
    for i in range(start, end):
        p = paras[i]
        text = p.text.strip()
        if not text:
            continue
        if _is_bullet(p):
            if current is not None:
                current.bullet_paras.append(p)
            continue
        if "|" in text:
            # New job heading: "Title | Company\tDates"
            left, _, right = text.partition("|")
            company_part, _, dates = right.partition("\t")
            current = ParsedExperience(
                title=left.strip(),
                company=company_part.strip(),
                dates=dates.strip(),
                heading_text=text,
            )
            experiences.append(current)
        # Non-heading, non-bullet lines (e.g. the location line) are ignored.
    return experiences


def collect_cv_skills(doc: Document) -> list[str]:
    """Pull the skills listed under KENNTNISSE, flattened across the category lines."""
    paras = doc.paragraphs
    start = None
    for i, p in enumerate(paras):
        if p.text.strip().upper() == "KENNTNISSE":
            start = i + 1
            break
    if start is None:
        return []
    skills: list[str] = []
    for i in range(start, len(paras)):
        t = paras[i].text.strip()
        if not t:
            continue
        if _SECTION_HEADING.match(t) and "|" not in t:
            break  # next section
        # Lines look like "Programming:  Python, Go, SQL". Take the part after ":".
        _, _, rest = t.partition(":")
        for token in re.split(r"[,/]", rest or t):
            token = token.strip()
            if 1 < len(token) <= 40:
                skills.append(token)
    # De-dupe, preserve order.
    seen = set()
    unique = []
    for s in skills:
        if s.lower() not in seen:
            seen.add(s.lower())
            unique.append(s)
    return unique


def apply_cv_bullets(experiences: list[ParsedExperience], new_bullets: list[list[str]]) -> None:
    """Write tailored bullets back, one list per experience, 1:1 by position.

    Requires each new list to match the count of existing bullet paragraphs, so
    no paragraphs are added or removed — the layout is untouched.
    """
    for exp, bullets in zip(experiences, new_bullets):
        if len(bullets) != len(exp.bullet_paras):
            raise ValueError(
                f"{exp.company}: got {len(bullets)} bullets for "
                f"{len(exp.bullet_paras)} paragraphs; counts must match"
            )
        for paragraph, text in zip(exp.bullet_paras, bullets):
            set_paragraph_text(paragraph, text)


# --- Anschreiben ---------------------------------------------------------


def _norm(text: str) -> str:
    """NFC-normalize for matching — .docx often stores umlauts decomposed."""
    return unicodedata.normalize("NFC", text)


_CLOSING = re.compile(r"freundlichen\s+Gr", re.IGNORECASE)


def _letter_cell(doc: Document) -> _Cell | None:
    """The merged cell holding the letter (date, subject, body, closing).

    The template lays the letter out in a table; the body cell is the one that
    contains the closing line. Merged cells appear once per grid slot, but the
    first hit is returned so duplicates are harmless.
    """
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                joined = _norm("\n".join(p.text for p in cell.paragraphs))
                if _CLOSING.search(joined):
                    return cell
    return None


@dataclass
class ParsedLetter:
    cell: _Cell
    subject_para: Paragraph | None
    body_paras: list[Paragraph]


def parse_anschreiben(doc: Document) -> ParsedLetter | None:
    """Locate the subject line and the body paragraphs (the pitch).

    Body = the non-empty paragraphs between the subject line and the closing
    ("Mit freundlichen Grüßen"). Header, date, closing, and name are left alone.
    """
    cell = _letter_cell(doc)
    if cell is None:
        return None

    paras = cell.paragraphs
    closing_idx = next(
        (i for i, p in enumerate(paras) if _CLOSING.search(_norm(p.text))),
        len(paras),
    )
    # The subject is the first substantial line; the date precedes it. Take the
    # longest of the first few non-empty lines before the body as the subject.
    non_empty = [(i, p) for i, p in enumerate(paras[:closing_idx]) if p.text.strip()]
    subject_para = None
    body_start = 0
    if non_empty:
        # Heuristic: date is short; subject is the next non-empty line.
        # Skip a leading date-like line.
        idx = 0
        if re.search(r"\d{4}", non_empty[0][1].text) and len(non_empty[0][1].text) < 25:
            idx = 1
        if idx < len(non_empty):
            subject_para = non_empty[idx][1]
            body_start = non_empty[idx][0] + 1

    body_paras = [
        p for i, p in enumerate(paras)
        if body_start <= i < closing_idx and p.text.strip()
    ]
    return ParsedLetter(cell=cell, subject_para=subject_para, body_paras=body_paras)


def apply_anschreiben(letter: ParsedLetter, body_paragraphs: list[str], subject: str | None) -> None:
    if subject and letter.subject_para is not None:
        set_paragraph_text(letter.subject_para, subject)
    if len(body_paragraphs) != len(letter.body_paras):
        raise ValueError(
            f"got {len(body_paragraphs)} body paragraphs for "
            f"{len(letter.body_paras)} slots; counts must match"
        )
    for paragraph, text in zip(letter.body_paras, body_paragraphs):
        set_paragraph_text(paragraph, text)
